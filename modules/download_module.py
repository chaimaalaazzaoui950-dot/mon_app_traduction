# modules/download_module.py

from docx import Document
from docx.shared import Pt
from fpdf import FPDF

# --- Fonctions spécifiques à chaque format ---

def save_txt(original_text, translated_text, source_lang, target_lang, filename):
    """Crée un fichier TXT formaté."""
    content = (
        f"--- Original ({source_lang.upper()}) ---\n"
        f"{original_text}\n\n"
        f"--- Traduction ({target_lang.upper()}) ---\n"
        f"{translated_text}"
    )
    with open(filename, "w", encoding="utf-8") as f:
        f.write(content)
    return filename

def save_docx(original_text, translated_text, source_lang, target_lang, filename):
    """Crée un fichier DOCX formaté."""
    doc = Document()
    
    # Titre pour le texte original
    p1_heading = doc.add_paragraph()
    p1_heading.add_run(f"Original ({source_lang.upper()})").bold = True
    
    # Contenu original
    doc.add_paragraph(original_text)
    
    # Espace
    doc.add_paragraph() 
    
    # Titre pour la traduction
    p2_heading = doc.add_paragraph()
    p2_heading.add_run(f"Traduction ({target_lang.upper()})").bold = True
    
    # Contenu traduit
    doc.add_paragraph(translated_text)
    
    doc.save(filename)
    return filename

def save_pdf(original_text, translated_text, source_lang, target_lang, filename):
    """Crée un fichier PDF formaté (gère les caractères spéciaux)."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    
    # FPDF gère mal l'UTF-8 par défaut. Il faut lui dire où trouver une police qui le supporte.
    # On va utiliser "DejaVu", une police libre qui gère la plupart des langues.
    # Vous devez avoir le fichier .ttf dans le même dossier ou spécifier le chemin.
    # Si vous n'avez pas cette police, Arial sera utilisé mais l'arabe/autres ne marcheront pas.
    try:
        pdf.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
        font_family = 'DejaVu'
    except RuntimeError:
        print("Police DejaVu non trouvée, utilisation d'Arial (limité).")
        font_family = 'Arial'

    # Titre pour le texte original
    pdf.set_font(font_family, 'B', 14)
    pdf.cell(0, 10, f"Original ({source_lang.upper()})", ln=True)
    
    # Contenu original
    pdf.set_font(font_family, '', 12)
    pdf.multi_cell(0, 10, original_text)
    
    # Espace
    pdf.ln(10)
    
    # Titre pour la traduction
    pdf.set_font(font_family, 'B', 14)
    pdf.cell(0, 10, f"Traduction ({target_lang.upper()})", ln=True)
    
    # Contenu traduit
    pdf.set_font(font_family, '', 12)
    pdf.multi_cell(0, 10, translated_text)
    
    pdf.output(filename)
    return filename

# --- Fonction principale appelée par l'application ---

def save_translation(original_text, translated_text, source_lang, target_lang, format, filename=None):
    """
    Dispatcher qui appelle la bonne fonction de sauvegarde en fonction du format.
    """
    if not filename:
        filename = f"traduction.{format}"

    args = {
        "original_text": original_text,
        "translated_text": translated_text,
        "source_lang": source_lang,
        "target_lang": target_lang,
        "filename": filename
    }

    if format == "txt":
        return save_txt(**args)
    elif format == "docx":
        return save_docx(**args)
    elif format == "pdf":
        return save_pdf(**args)
    else:
        raise ValueError("Format de sauvegarde non supporté.")

